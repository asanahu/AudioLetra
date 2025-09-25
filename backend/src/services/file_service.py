"""
File generation service for creating downloadable documents
Supports .txt, .docx, and .pdf formats
"""
import os
import tempfile
from typing import Dict, Any, Optional, Tuple
from io import BytesIO
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from src.config import Config


class FileService:
    """Service for generating and managing downloadable files"""
    
    def __init__(self, config: Optional[Config] = None):
        self.config = config or Config()
        self.temp_dir = tempfile.gettempdir()
    
    def generate_txt_file(self, content: str, filename: str) -> Tuple[str, str]:
        """Generate a TXT file"""
        file_path = os.path.join(self.temp_dir, filename)
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return file_path, 'text/plain'
        except Exception as e:
            raise Exception(f"Failed to generate TXT file: {str(e)}")
    
    def generate_docx_file(self, content: str, filename: str, title: str = None) -> Tuple[str, str]:
        """Generate a DOCX file"""
        file_path = os.path.join(self.temp_dir, filename)
        
        try:
            doc = Document()
            
            # Add title if provided
            if title:
                title_paragraph = doc.add_heading(title, level=1)
                title_paragraph.alignment = 1  # Center alignment
            
            # Add content
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Add footer with generation info
            footer_section = doc.sections[0]
            footer = footer_section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = f"Generado por AudioLetra - {self._get_current_timestamp()}"
            
            doc.save(file_path)
            return file_path, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
        except Exception as e:
            raise Exception(f"Failed to generate DOCX file: {str(e)}")
    
    def generate_pdf_file(self, content: str, filename: str, title: str = None) -> Tuple[str, str]:
        """Generate a PDF file"""
        file_path = os.path.join(self.temp_dir, filename)
        
        try:
            doc = SimpleDocTemplate(file_path, pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Container for the 'Flowable' objects
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Add title if provided
            if title:
                title_style = styles['Title']
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 12))
            
            # Add content
            normal_style = styles['Normal']
            paragraphs = content.split('\n\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Handle line breaks within paragraphs
                    paragraph_lines = paragraph.split('\n')
                    formatted_paragraph = '<br/>'.join(paragraph_lines)
                    story.append(Paragraph(formatted_paragraph, normal_style))
                    story.append(Spacer(1, 12))
            
            # Add footer
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=8,
                textColor='gray'
            )
            story.append(Spacer(1, 24))
            story.append(Paragraph(f"Generado por AudioLetra - {self._get_current_timestamp()}", footer_style))
            
            # Build PDF
            doc.build(story)
            return file_path, 'application/pdf'
            
        except Exception as e:
            raise Exception(f"Failed to generate PDF file: {str(e)}")
    
    def generate_file(self, content: str, format_type: str, base_filename: str, title: str = None) -> Tuple[str, str]:
        """Generate a file in the specified format"""
        # Validate format
        supported_formats = ['txt', 'docx', 'pdf']
        if format_type.lower() not in supported_formats:
            raise ValueError(f"Unsupported format: {format_type}. Supported formats: {', '.join(supported_formats)}")
        
        # Generate filename with extension
        if not base_filename.endswith(f'.{format_type}'):
            filename = f"{base_filename}.{format_type}"
        else:
            filename = base_filename
        
        # Sanitize filename
        filename = self._sanitize_filename(filename)
        
        # Generate file based on format
        format_type = format_type.lower()
        
        if format_type == 'txt':
            return self.generate_txt_file(content, filename)
        elif format_type == 'docx':
            return self.generate_docx_file(content, filename, title)
        elif format_type == 'pdf':
            return self.generate_pdf_file(content, filename, title)
    
    def read_file_content(self, file_path: str) -> bytes:
        """Read file content as bytes"""
        try:
            with open(file_path, 'rb') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Failed to read file: {str(e)}")
    
    def delete_file(self, file_path: str) -> bool:
        """Delete a file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception:
            return False
    
    def get_content_type(self, format_type: str) -> str:
        """Get content type for a format"""
        content_types = {
            'txt': 'text/plain',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'pdf': 'application/pdf'
        }
        return content_types.get(format_type.lower(), 'application/octet-stream')
    
    def cleanup_temp_files(self, older_than_hours: int = 1):
        """Clean up temporary files older than specified hours"""
        import time
        
        cutoff_time = time.time() - (older_than_hours * 3600)
        cleaned_count = 0
        
        try:
            for filename in os.listdir(self.temp_dir):
                if filename.startswith('audioletra_'):
                    file_path = os.path.join(self.temp_dir, filename)
                    if os.path.isfile(file_path):
                        file_mtime = os.path.getmtime(file_path)
                        if file_mtime < cutoff_time:
                            try:
                                os.remove(file_path)
                                cleaned_count += 1
                            except Exception:
                                pass  # Ignore errors for individual files
        except Exception:
            pass  # Ignore errors in directory listing
        
        return cleaned_count
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent path traversal and invalid characters"""
        # Remove path separators and invalid characters
        invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Limit length
        if len(filename) > 100:
            name, ext = os.path.splitext(filename)
            filename = name[:100-len(ext)] + ext
        
        # Add prefix to avoid conflicts
        if not filename.startswith('audioletra_'):
            name, ext = os.path.splitext(filename)
            filename = f"audioletra_{name}{ext}"
        
        return filename
    
    def _get_current_timestamp(self) -> str:
        """Get current timestamp as formatted string"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get information about a file"""
        try:
            if not os.path.exists(file_path):
                return {"exists": False}
            
            stat = os.stat(file_path)
            
            return {
                "exists": True,
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "filename": os.path.basename(file_path),
                "extension": os.path.splitext(file_path)[1].lower(),
                "path": file_path
            }
        except Exception as e:
            return {"exists": False, "error": str(e)}


# Global service instance
_file_service: Optional[FileService] = None


def get_file_service(config: Optional[Config] = None) -> FileService:
    """Get or create the global file service instance"""
    global _file_service
    
    if _file_service is None:
        _file_service = FileService(config)
    
    return _file_service


def reset_file_service():
    """Reset the global file service instance"""
    global _file_service
    _file_service = None
